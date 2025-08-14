import uuid
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib
import io
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import TokenTextSplitter

class DocumentProcessor:
    def __init__(self, embeddings_client, vector_store):
        self.embeddings_client = embeddings_client
        self.vector_store = vector_store
        self.logger = logging.getLogger(__name__)
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_stream = io.BytesIO(file_content)
            reader = PdfReader(pdf_stream)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
                    
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF: {e}")
            return ""
        
        return text.strip()
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc_stream = io.BytesIO(file_content)
            doc = Document(doc_stream)
            
            # Extract text from paragraphs and tables
            text_parts = []
            
            # Get paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Get table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            return "\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX: {e}")
            return ""

    def chunk_text(self, text: str, max_tokens: int = 800, overlap: int = 200) -> List[str]:
        """Split text into chunks using token-based splitting"""
        try:
            splitter = TokenTextSplitter(
                chunk_size=max_tokens,
                chunk_overlap=overlap
            )
            chunks = splitter.split_text(text)
            
            # Filter out very small chunks
            filtered_chunks = [chunk for chunk in chunks if len(chunk.strip()) > 50]
            
            return filtered_chunks
            
        except Exception as e:
            self.logger.error(f"Failed to chunk text: {e}")
            # Fallback to simple chunking
            return self._simple_chunk_text(text, max_tokens * 4, overlap * 4)  # Rough char estimate
    
    def _simple_chunk_text(self, text: str, chunk_size: int = 3200, overlap: int = 800) -> List[str]:
        """Fallback simple text chunking by characters"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                sentence_end = text.rfind('.', start + chunk_size - 200, end)
                if sentence_end != -1:
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
                
        return chunks
    
    async def process_document(self, file_content: bytes, filename: str, 
                             user_id: str, file_type: str, document_id: str = None) -> Dict[str, Any]:
        """
        Main processing pipeline
        """
        try:
            # Step 1: Extract text based on file type
            self.logger.info(f"Starting document processing for {filename} (type: {file_type})")
        
            if file_type.lower() == 'pdf':
                text = self.extract_text_from_pdf(file_content)
            elif file_type.lower() in ['docx', 'doc']:
                text = self.extract_text_from_docx(file_content)
            elif file_type.lower() == 'txt':
                try:
                    text = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try other encodings
                    for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                        try:
                            text = file_content.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        raise ValueError("Unable to decode text file")
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text or len(text.strip()) < 10:
                raise ValueError("No meaningful text content found in document")
            
            # Step 2: Generate document ID and metadata
            doc_id = document_id if document_id else str(uuid.uuid4())
            file_hash = hashlib.md5(file_content).hexdigest()
            
            # Step 3: Chunk the text
            chunks = self.chunk_text(text)
            
            if not chunks:
                raise ValueError("Failed to create chunks from document text")
            
            # Step 4: Process each chunk (embed and store)
            stored_chunks = []
            failed_chunks = 0
            
            for i, chunk in enumerate(chunks):
                try:
                    chunk_id = f"{doc_id}_chunk_{i}"
                    
                    # Generate embedding for the chunk
                    embedding = self.embeddings_client.embed_query(chunk)  
                    
                    if not embedding:
                        failed_chunks += 1
                        continue
                    
                    # Create chunk metadata
                    chunk_metadata = {
                        "document_id": doc_id,
                        "chunk_index": i,
                        "user_id": user_id,
                        "filename": filename,
                        "file_type": file_type,
                        "file_hash": file_hash,
                        "chunk_text": chunk,
                        "chunk_length": len(chunk),
                        "timestamp": datetime.now().isoformat()
                    }

                    # Store in vector database
                    success = self.vector_store.store_conversation(
                        user_id=f"{user_id}_docs",  
                        conversation_text=chunk,
                        embedding=embedding,
                        metadata=chunk_metadata
                    )
                    
                    if success:
                        stored_chunks.append({
                            "chunk_id": chunk_id,
                            "preview": chunk[:100] + "..." if len(chunk) > 100 else chunk
                        })
                    else:
                        failed_chunks += 1
                        
                except Exception as e:
                    failed_chunks += 1
                    continue
            
            if not stored_chunks:
                raise ValueError("Failed to store any chunks from the document")
            
            success_rate = len(stored_chunks) / len(chunks)
            self.logger.info(f"Document processing completed: {len(stored_chunks)}/{len(chunks)} chunks stored successfully")
            
            return {
                "document_id": doc_id,
                "filename": filename,
                "file_type": file_type,
                "file_hash": file_hash,
                "total_chunks": len(chunks),
                "stored_chunks_count": len(stored_chunks),
                "failed_chunks": failed_chunks,
                "success_rate": round(success_rate, 2),
                "text_length": len(text),
                "stored_chunks": stored_chunks[:5],  
                "status": "success" if success_rate > 0.5 else "partial_success"
            }
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {e}")
            return {
                "error": str(e),
                "filename": filename,
                "status": "failed"
            }

class DocumentRetriever:
    def __init__(self, embeddings_client, vector_store):
        self.embeddings_client = embeddings_client
        self.vector_store = vector_store
        self.logger = logging.getLogger(__name__)
    
    async def search_documents(self, query: str, user_id: str, 
                             top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for relevant document chunks"""
        try:
            # Generate embedding for query
            query_embedding = self.embeddings_client.embed_query(query)  # Remove await - sync method
            
            if not query_embedding:
                self.logger.warning(f"Failed to generate embedding for query: {query}")
                return []

            # Search in documents namespace
            similar_chunks = self.vector_store.similarity_search(
                user_id=f"{user_id}_docs",  # Use documents namespace
                query_embedding=query_embedding,
                top_k=top_k
            )
            
            # Format results
            results = []
            for chunk_data in similar_chunks:
                metadata = chunk_data.get('metadata', {})
                results.append({
                    "document_id": metadata.get('document_id'),
                    "filename": metadata.get('filename'),
                    "chunk_index": metadata.get('chunk_index'),
                    "content": metadata.get('chunk_text', '')[:500],  # Limit content length
                    "score": chunk_data.get('score', 0),
                    "file_type": metadata.get('file_type'),
                    "timestamp": metadata.get('timestamp')
                })
            
            self.logger.info(f"Document search returned {len(results)} results for query: {query[:50]}...")
            return results
            
        except Exception as e:
            self.logger.error(f"Error searching documents: {e}")
            return []
    async def search_specific_documents(self, query: str, user_id: str, 
                                  document_ids: List[str], top_k: int = 5) -> List[Dict[str, Any]]:
        try:
            # Generate embedding for query
            query_embedding = self.embeddings_client.embed_query(query)

            if not query_embedding:
                self.logger.warning(f"Failed to generate embedding for query: {query}")
                return []


            # Search in documents namespace with document_id filter
            similar_chunks = self.vector_store.similarity_search_with_filter(
                user_id=f"{user_id}_docs",
                query_embedding=query_embedding,
                top_k=top_k,
                filter_condition={"document_id": {"$in": document_ids}}  # Filter by document IDs
            )
        
            # Format results
            results = []
            for chunk_data in similar_chunks:
                metadata = chunk_data.get('metadata', {})
                if metadata.get('document_id') in document_ids:  # Double-check
                    results.append({
                        "document_id": metadata.get('document_id'),
                        "filename": metadata.get('filename'),
                        "chunk_index": metadata.get('chunk_index'),
                        "content": metadata.get('chunk_text', '')[:500],
                        "score": chunk_data.get('score', 0),
                        "file_type": metadata.get('file_type'),
                        "timestamp": metadata.get('timestamp')
                    })

            self.logger.info(f"Specific document search returned {len(results)} results")
            return results
        
        except Exception as e:
            self.logger.error(f"Error searching specific documents: {e}")
            return []
        
